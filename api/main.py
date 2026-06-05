import io
import json
import logging
import os
import sys
import threading
import time
import traceback
import uvicorn
from fastapi import Depends, FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, StreamingResponse
from langchain_core.prompts import ChatPromptTemplate
from logging.handlers import RotatingFileHandler
from pydantic import BaseModel

from src.utils.DatabaseHandler import DatabaseHandler, DATABASE_DIR
from src.utils.LLMHandler import LLMHandler
from src.utils.NotePersistenceHandler import NotePersistenceHandler
from src.utils.SummaryHandler import SummaryHandler

_USER_DATA_FILE = "data/user_data.json"
_NOTES_FILE = "data/editor_notes.txt"
_LOG_FILE = "data/operations.log"
_LOG_MAX_BYTES = 5 * 1024 * 1024
_LOG_BACKUP_COUNT = 3
_OP_LOGGER_NAME = "dandd.operations"

_LLM_CALL_TIMEOUT_SECONDS = 180
_CHAT_MAX_PREDICT = 2048


def _get_op_logger() -> logging.Logger:
    return logging.getLogger(_OP_LOGGER_NAME)


def _configure_op_logger(log_file: str) -> None:
    logger = logging.getLogger(_OP_LOGGER_NAME)
    for h in list(logger.handlers):
        if isinstance(h, RotatingFileHandler):
            h.close()
            logger.removeHandler(h)
    log_dir = os.path.dirname(log_file)
    if log_dir:
        os.makedirs(log_dir, exist_ok=True)
    fh = RotatingFileHandler(log_file, maxBytes=_LOG_MAX_BYTES, backupCount=_LOG_BACKUP_COUNT)
    fh.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
    logger.addHandler(fh)
    logger.setLevel(logging.DEBUG)

_singleton_lock = threading.Lock()


def _invoke_with_timeout(invoke_callable, *args, **kwargs):
    result = [None]
    exc = [None]
    done = threading.Event()

    def _run():
        try:
            result[0] = invoke_callable(*args, **kwargs)
        except Exception as e:
            exc[0] = e
        finally:
            done.set()

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    t.join(timeout=_LLM_CALL_TIMEOUT_SECONDS)

    if not done.is_set():
        frames = sys._current_frames()
        frame = frames.get(t.ident)
        stack_str = "".join(traceback.format_stack(frame)) if frame is not None else "(frame unavailable)"
        _get_op_logger().error(
            f"LLM call timed out after {_LLM_CALL_TIMEOUT_SECONDS}s\nStuck thread stack:\n{stack_str}"
        )
        raise TimeoutError(f"LLM call exceeded {_LLM_CALL_TIMEOUT_SECONDS}s timeout")
    if exc[0] is not None:
        raise exc[0]
    return result[0]


_db_singleton: DatabaseHandler | None = None
_llm_singleton: LLMHandler | None = None

_CHAT_PROMPT = ChatPromptTemplate.from_messages([
    (
        "system",
        "You are an expert in answering questions about a TTRPG campaign described in provided documents. "
        "The provided documents describe a campaign where the party members (player characters) are {partymembers}. "
        "Here are the relevant documents from {notetaker}'s perspective (could be in first person or third person): {notes}"
        "\n Base your answers only off of the provided documents. "
        "Do not use extraneous information to answer the question. "
        "Do not provide references to the documents.",
    ),
    ("user", "{question}"),
])


class _TimedLLMHandler:
    def __init__(self, handler):
        self._handler = handler

    def __getattr__(self, name):
        return getattr(self._handler, name)

    def invoke_model(self, prompt, mappings):
        return _invoke_with_timeout(self._handler.invoke_model, prompt, mappings)


def get_llm_handler() -> LLMHandler:
    global _llm_singleton
    if _llm_singleton is None:
        with _singleton_lock:
            if _llm_singleton is None:
                _llm_singleton = LLMHandler()
    return _llm_singleton


def get_db_handler() -> DatabaseHandler:
    global _db_singleton
    if _db_singleton is None:
        with _singleton_lock:
            if _db_singleton is None:
                h = DatabaseHandler()
                h.create_retrival_artifacts(DATABASE_DIR)
                _db_singleton = h
    return _db_singleton


def get_persistence_handler() -> NotePersistenceHandler:
    return NotePersistenceHandler()


def get_summary_handler() -> SummaryHandler:
    return SummaryHandler(_TimedLLMHandler(get_llm_handler()))


def _read_file_as_text(file_path: str) -> str:
    """Read a file and return its content as plain text."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        from docx import Document as DocxReader
        with open(file_path, "rb") as f:
            doc = DocxReader(io.BytesIO(f.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _sse_event(payload: dict) -> str:
    return f"data: {json.dumps(payload)}\n\n"


class UploadNotesRequest(BaseModel):
    file_path: str


class ChatRequest(BaseModel):
    question: str
    model: str
    temperature: float


class SummaryGenerateRequest(BaseModel):
    model: str
    party_members: list[str] = []
    temperature: float = 0.7


class NotesRequest(BaseModel):
    content: str


class PartyRequest(BaseModel):
    party_members: list[dict] = []


def create_app() -> FastAPI:
    _configure_op_logger(_LOG_FILE)
    application = FastAPI()
    _heavy_lock = threading.Lock()

    def _guarded_stream(inner, op_name: str = "operation"):
        """Wraps an SSE generator with the global heavy-op lock and operation logging."""
        def guarded():
            logger = _get_op_logger()
            if not _heavy_lock.acquire(blocking=False):
                logger.warning(f"[{op_name}] busy rejection: lock held by another operation")
                yield _sse_event({
                    "done": True,
                    "error": True,
                    "message": "Backend is busy. Wait for the current operation to finish.",
                })
                return
            logger.info(f"[{op_name}] lock acquired")
            t_start = time.monotonic()
            logger.info(f"[{op_name}] start")
            try:
                yield from inner()
            finally:
                duration = time.monotonic() - t_start
                logger.info(f"[{op_name}] end, duration={duration:.3f}s")
                _heavy_lock.release()
                logger.info(f"[{op_name}] lock released")
        return guarded

    application.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_methods=["*"],
        allow_headers=["*"],
    )

    @application.get("/health")
    def health():
        return {"status": "ok"}

    @application.get("/status")
    def status(persistence: NotePersistenceHandler = Depends(get_persistence_handler)):
        return {"has_notes": persistence.has_notes()}

    @application.get("/models")
    def models(handler: LLMHandler = Depends(get_llm_handler)):
        return [m.model for m in handler.get_available_models()]

    @application.post("/upload-notes")
    def upload_notes(
        body: UploadNotesRequest,
        db: DatabaseHandler = Depends(get_db_handler),
        persistence: NotePersistenceHandler = Depends(get_persistence_handler),
    ):
        def event_stream():
            try:
                db.clear_database(DATABASE_DIR)
                db.create_retrival_artifacts(DATABASE_DIR)

                class _FileWrapper:
                    def __init__(self, path: str):
                        self.name = os.path.basename(path)
                        self._path = path

                    def read(self):
                        with open(self._path, "rb") as f:
                            return f.read()

                    def getvalue(self):
                        return self.read()

                doc = _FileWrapper(body.file_path)
                for pct in db.generate_database(doc, DATABASE_DIR):
                    yield _sse_event({
                        "done": False,
                        "progress": int(pct),
                        "message": f"Processing… {int(pct)}%",
                    })
                if db.last_processed_df is not None and not db.last_processed_df.empty:
                    persistence.persist(db.last_processed_df, _read_file_as_text(body.file_path))
                yield _sse_event({"done": True, "progress": 100})
            except Exception as exc:
                _get_op_logger().error("[upload-notes] error: %s", exc, exc_info=True)
                yield _sse_event({"done": True, "error": True, "message": str(exc)})

        return StreamingResponse(_guarded_stream(event_stream, "upload-notes")(), media_type="text/event-stream")

    @application.post("/chat")
    def chat(
        body: ChatRequest,
        llm: LLMHandler = Depends(get_llm_handler),
        db: DatabaseHandler = Depends(get_db_handler),
    ):
        def event_stream():
            try:
                yield _sse_event({"done": False, "progress": 10, "message": "Retrieving relevant notes…"})

                db.create_retrival_artifacts(DATABASE_DIR)
                notes = db.retrieve_notes(body.question)

                yield _sse_event({"done": False, "progress": 50, "message": "Generating response…"})

                party_members: list[dict] = []
                note_taker = "Unknown"
                try:
                    if os.path.isfile(_USER_DATA_FILE):
                        with open(_USER_DATA_FILE, "r") as f:
                            user_data = json.load(f)
                        party_members = user_data.get("party_members", [])
                        takers = [m.get("name", "") for m in party_members if m.get("note_taker")]
                        if takers:
                            note_taker = takers[0]
                except Exception:
                    pass

                member_names = [m.get("name", "") for m in party_members if m.get("name")]
                if len(member_names) > 1:
                    formatted_members = ", ".join(member_names[:-1]) + ", and " + member_names[-1]
                elif member_names:
                    formatted_members = member_names[0]
                else:
                    formatted_members = "the party"

                if notes:
                    llm.load_model(body.model, body.temperature, disable_thinking=True, num_predict=_CHAT_MAX_PREDICT)
                    answer = _invoke_with_timeout(
                        llm.invoke_model,
                        _CHAT_PROMPT,
                        {
                            "question": body.question,
                            "partymembers": formatted_members,
                            "notes": notes,
                            "notetaker": note_taker,
                        },
                    )
                else:
                    answer = (
                        "Could not find any relevant journal entries for your query. "
                        "It could be that there is not any relevant information regarding your query in the notes, "
                        "the question needs to be reworded, or spelling needs to be reviewed."
                    )

                sources = [
                    {"content": doc.page_content, "date": doc.metadata.get("Date", "Unknown")}
                    for doc in notes
                ]
                yield _sse_event({"done": True, "answer": answer, "sources": sources})
            except Exception as exc:
                _get_op_logger().error("[chat] error: %s", exc, exc_info=True)
                yield _sse_event({"done": True, "error": True, "message": str(exc)})

        return StreamingResponse(_guarded_stream(event_stream, "chat")(), media_type="text/event-stream")

    @application.get("/summary")
    def summary_get(summary: SummaryHandler = Depends(get_summary_handler)):
        data = summary.get_saved_summary()
        return data if data is not None else {}

    @application.get("/notes")
    def notes_get():
        if os.path.isfile(_NOTES_FILE):
            with open(_NOTES_FILE, "r", encoding="utf-8") as f:
                return {"content": f.read()}
        return {"content": ""}

    @application.post("/notes")
    def notes_post(body: NotesRequest):
        os.makedirs(os.path.dirname(_NOTES_FILE), exist_ok=True)
        with open(_NOTES_FILE, "w", encoding="utf-8") as f:
            f.write(body.content)
        return {"status": "ok"}

    @application.post("/notes/vectorize")
    def notes_vectorize(
        body: NotesRequest,
        db: DatabaseHandler = Depends(get_db_handler),
        persistence: NotePersistenceHandler = Depends(get_persistence_handler),
    ):
        def event_stream():
            try:
                db.create_retrival_artifacts(DATABASE_DIR)

                class _TextWrapper:
                    name = "notes.txt"

                    def __init__(self, content: str):
                        self._bytes = content.encode("utf-8")

                    def read(self):
                        return self._bytes

                    def getvalue(self):
                        return self._bytes

                wrapper = _TextWrapper(body.content)
                for pct in db.generate_database(wrapper, DATABASE_DIR):
                    yield _sse_event({
                        "done": False,
                        "progress": int(pct),
                        "message": f"Vectorizing… {int(pct)}%",
                    })
                if db.last_processed_df is not None and not db.last_processed_df.empty:
                    persistence.persist(db.last_processed_df, body.content)
                yield _sse_event({"done": True, "progress": 100})
            except Exception as exc:
                _get_op_logger().error("[notes-vectorize] error: %s", exc, exc_info=True)
                yield _sse_event({"done": True, "error": True, "message": str(exc)})

        return StreamingResponse(_guarded_stream(event_stream, "notes-vectorize")(), media_type="text/event-stream")

    @application.post("/summary/generate")
    def summary_generate(
        body: SummaryGenerateRequest,
        summary: SummaryHandler = Depends(get_summary_handler),
    ):
        def event_stream():
            try:
                for is_done, progress, text in summary.generate_summary_streaming(
                    body.model, body.party_members, temperature=body.temperature
                ):
                    if is_done:
                        yield _sse_event({"done": True, "progress": 100})
                    else:
                        yield _sse_event({"done": False, "progress": progress, "message": text})
            except Exception as exc:
                _get_op_logger().error("[summary-generate] error: %s", exc, exc_info=True)
                yield _sse_event({"done": True, "error": True, "message": str(exc)})

        return StreamingResponse(_guarded_stream(event_stream, "summary-generate")(), media_type="text/event-stream")

    @application.get("/party")
    def party_get():
        if os.path.isfile(_USER_DATA_FILE):
            with open(_USER_DATA_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            return {"party_members": data.get("party_members", [])}
        return {"party_members": []}

    @application.post("/party")
    def party_post(body: PartyRequest):
        data = {}
        if os.path.isfile(_USER_DATA_FILE):
            with open(_USER_DATA_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
        data["party_members"] = body.party_members
        os.makedirs(os.path.dirname(_USER_DATA_FILE), exist_ok=True)
        with open(_USER_DATA_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f)
        return {"status": "ok"}

    @application.get("/notes/export/txt")
    def notes_export_txt():
        if not os.path.isfile(_NOTES_FILE):
            from fastapi import HTTPException
            raise HTTPException(status_code=404, detail="No notes file found.")
        with open(_NOTES_FILE, "rb") as f:
            content = f.read()
        return Response(
            content=content,
            media_type="text/plain",
            headers={"Content-Disposition": "attachment; filename=\"editor_notes.txt\""},
        )

    @application.get("/notes/export/docx")
    def notes_export_docx():
        if not os.path.isfile(_NOTES_FILE):
            from fastapi import HTTPException
            raise HTTPException(status_code=404, detail="No notes file found.")
        from docx import Document
        with open(_NOTES_FILE, "r", encoding="utf-8") as f:
            text = f.read()
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=\"editor_notes.docx\""},
        )

    return application


app = create_app()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("api.main:app", host="127.0.0.1", port=port, reload=False)
