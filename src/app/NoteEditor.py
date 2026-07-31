"""Note editor page for creating and editing campaign notes."""

import io
import json

import pandas as pd
import streamlit as st
from docx import Document as _DocxDocument
from html.parser import HTMLParser
from streamlit_lottie import st_lottie

from ..utils import DatabaseHandler
from ..utils.SummaryHandler import SummaryHandler
from ..utils.TextEditorHandler import TextEditorHandler

_RAW_NOTES_KEY = SummaryHandler.RAW_NOTES_KEY
_SUMMARY_KEY = SummaryHandler.SUMMARY_KEY


# ---------------------------------------------------------------------------
# Module-level utilities — pure functions with no Streamlit dependency
# ---------------------------------------------------------------------------

class _HTMLTextExtractor(HTMLParser):
    _BLOCK_TAGS = frozenset({"p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"})

    def __init__(self):
        super().__init__()
        self._parts = []

    def handle_starttag(self, tag, attrs):
        if tag in self._BLOCK_TAGS:
            self._parts.append("\n")

    def handle_data(self, data):
        self._parts.append(data)

    def get_text(self):
        return "".join(self._parts).strip()


def strip_html(html_content: str) -> str:
    """Return plain text with HTML tags removed."""
    extractor = _HTMLTextExtractor()
    extractor.feed(html_content or "")
    return extractor.get_text()


def raw_notes_to_text(raw_notes_json: str) -> str:
    """Convert the raw-notes DataFrame JSON string (from session state) to plain text."""
    if not raw_notes_json:
        return ""
    try:
        df = pd.read_json(io.StringIO(raw_notes_json))
    except Exception:
        return ""
    parts = []
    for _, row in df.iterrows():
        date = str(row.get("Date", "") or "")
        title = str(row.get("Title", "") or "")
        contents = str(row.get("Contents", "") or "")
        header = date if date.lower() not in ("", "unknown date", "nan") else title
        if header:
            parts.append(header)
        if contents:
            parts.append(contents)
        parts.append("")
    return "\n".join(parts).strip()


def build_txt_content(content: str) -> str:
    """Return plain text content for TXT export."""
    return content or ""


def build_docx_bytes(content: str) -> bytes:
    """Convert plain text to DOCX bytes for download."""
    doc = _DocxDocument()
    for line in (content or "").splitlines():
        stripped = line.strip()
        if stripped:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _EditorDocument:
    """Wraps plain text as a file-like object compatible with DatabaseHandler."""

    def __init__(self, text_content: str):
        self.name = "editor_notes.txt"
        self._data = text_content.encode("utf-8")

    def getvalue(self) -> bytes:
        return self._data

    def read(self) -> bytes:
        return self._data


# ---------------------------------------------------------------------------
# NoteEditor page class
# ---------------------------------------------------------------------------

class NoteEditor:
    def __init__(self):
        if "databasehandler" not in st.session_state:
            st.session_state.databasehandler = DatabaseHandler.DatabaseHandler()
        self.databasehandler = st.session_state.databasehandler
        self._editor = TextEditorHandler()

    def __init_state_variables(self):
        if 'is_processing' not in st.session_state:
            st.session_state.is_processing = False

        if "editor_content" not in st.session_state:
            st.session_state.editor_content = raw_notes_to_text(st.session_state.get(_RAW_NOTES_KEY))
        if "editor_key" not in st.session_state:
            st.session_state.editor_key = 0

        # Streamlit drops widget-keyed state when the widget is not rendered for
        # a full run (e.g. on page switch). Keep the preference in a separate
        # (non-widget) session key and re-seed the toggle from it whenever the
        # widget key is absent so the dark mode choice survives a page switch.
        if "editor_dark_mode" not in st.session_state:
            st.session_state.editor_dark_mode = bool(st.session_state.get("editor_dark_mode_pref", False))

    def __persist_dark_mode(self):
        st.session_state.editor_dark_mode_pref = bool(st.session_state.editor_dark_mode)

    def __import_uploaded_notes(self):
        text = raw_notes_to_text(st.session_state.get(_RAW_NOTES_KEY))
        if text:
            st.session_state.editor_content = text
            st.session_state.editor_key += 1

    def __vectorize_notes(self):
        content = strip_html(st.session_state.editor_content)
        if not content or not content.strip():
            st.error("No content in the editor to vectorize.")
            return

        openai_key = st.session_state.get("openai_api_key")
        if not openai_key:
            st.error("Enter your OpenAI API key in Model Options to vectorize notes.")
            st.session_state.is_processing = False
            return

        self.databasehandler.clear_database(DatabaseHandler.DATABASE_DIR)
        st.session_state.pop(_RAW_NOTES_KEY, None)
        st.session_state.pop(_SUMMARY_KEY, None)

        self.databasehandler.create_retrival_artifacts(DatabaseHandler.DATABASE_DIR, openai_key)

        doc = _EditorDocument(content)

        try:
            with open("assets/Magical_Effect_Loading.json", "r", errors="ignore") as f:
                magic_loader = json.load(f)
        except Exception:
            magic_loader = None

        animation_slot = st.empty()
        progress_slot = st.empty()

        if magic_loader:
            with animation_slot.container():
                st_lottie(magic_loader, height=200, key="editor_vectorize_spinner")

        progress_bar = progress_slot.progress(0, text="Vectorizing notes...")
        gen = self.databasehandler.generate_database(doc, DatabaseHandler.DATABASE_DIR)
        return_code = None

        while True:
            try:
                progress = next(gen)
                progress_bar.progress(progress / 100, text=f"Vectorizing notes... {progress:.1f}%")
            except StopIteration as e:
                return_code = e.value
                break

        animation_slot.empty()
        progress_slot.empty()

        if return_code:
            if self.databasehandler.last_processed_df is not None:
                st.session_state[_RAW_NOTES_KEY] = self.databasehandler.last_processed_df.to_json()
            st.toast("📜 Notes vectorized successfully!", icon="🧙‍♂️")
        else:
            st.error("Vectorization failed. Ensure the editor contains valid content.")
        st.session_state.is_processing = False

    def __render_sidebar(self):
        with st.sidebar:
            st.header("📜 Note Options")

            processing = st.session_state.is_processing
            raw_notes_exist = bool(st.session_state.get(_RAW_NOTES_KEY))
            if st.button(
                "📥 Load from Uploaded Notes",
                disabled=not raw_notes_exist or processing,
                help="No uploaded notes found." if not raw_notes_exist else "Replace editor content with notes from the last file upload.",
                use_container_width=True,
            ):
                self.__import_uploaded_notes()
                st.rerun()

            st.divider()

            st.subheader("🎨 Appearance")
            st.toggle(
                "🌙 Dark mode",
                key="editor_dark_mode",
                on_change=self.__persist_dark_mode,
                disabled=processing,
                help="Switch the note editor to a dark colour scheme.",
            )

            st.divider()

            st.subheader("🧠 Vectorize")
            has_content = bool(strip_html(st.session_state.get("editor_content", "")).strip())
            has_key = bool(st.session_state.get("openai_api_key"))
            vectorize_ready = has_content and has_key
            if not has_content:
                vectorize_help = "No content to vectorize."
            elif not has_key:
                vectorize_help = "Enter your OpenAI API key in Model Options to vectorize notes."
            else:
                vectorize_help = "Clear the database and re-vectorize the current editor notes."
            if st.button(
                "⚡ Vectorize Notes",
                type="primary",
                use_container_width=True,
                disabled=not vectorize_ready or processing,
                help=vectorize_help,
            ):
                st.session_state._do_vectorize = True
                st.session_state.is_processing = True
                st.rerun()

            st.divider()

            st.subheader("📤 Export")
            content = strip_html(st.session_state.get("editor_content", ""))
            st.download_button(
                label="📄 Export as TXT",
                data=build_txt_content(content),
                file_name="campaign_notes.txt",
                mime="text/plain",
                use_container_width=True,
                disabled=processing,
            )
            st.download_button(
                label="📝 Export as DOCX",
                data=build_docx_bytes(content),
                file_name="campaign_notes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                disabled=processing,
            )

    def __render_editor(self):
        new_content = self._editor.render(
            key=f"note_editor_{st.session_state.editor_key}",
            initial_value=st.session_state.get("editor_content", ""),
            dark_mode=st.session_state.get("editor_dark_mode", False),
        )
        if new_content != st.session_state.editor_content:
            st.session_state.editor_content = new_content

    def run(self):
        st.title("📝 Note Editor")
        st.info(
            "Create and edit your campaign notes here. "
            "Notes are kept for this session only and are cleared when the app is refreshed or restarted."
        )

        self.__init_state_variables()
        self.__render_sidebar()

        if st.session_state.pop("_do_vectorize", False):
            self.__vectorize_notes()
            st.rerun()

        self.__render_editor()
